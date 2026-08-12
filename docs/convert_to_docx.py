"""Convert THESIS_REPORT.md to THESIS_REPORT.docx using python-docx."""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def add_paragraph(doc, text, bold=False, italic=False, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def process_md_file(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                in_code_block = False
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                # Add shading
                from docx.oxml.ns import qn
                shading = run._element.get_or_add_rPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto',
                    qn('w:fill'): 'F5F5F5'
                })
                shading.append(shd)
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            if all(re.match(r'^[-:]+$', c) for c in cells):
                # separator row, skip
                i += 1
                continue
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            # Check if next line is not a table row
            if i + 1 >= len(lines) or '|' not in lines[i + 1] or not lines[i + 1].strip().startswith('|'):
                in_table = False
                if table_rows:
                    headers = table_rows[0]
                    rows = table_rows[1:]
                    # Pad rows to match header length
                    max_cols = len(headers)
                    for r in rows:
                        while len(r) < max_cols:
                            r.append('')
                    add_table(doc, headers, rows)
                    doc.add_paragraph()  # spacing after table
            i += 1
            continue

        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            if text and not text.startswith('Cybercrime Guidance Agentic AI'):
                add_heading(doc, text, 0)
            elif text:
                # Title page
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(24)
                run.font.color.rgb = RGBColor(0x0e, 0x44, 0x6a)
            i += 1
            continue

        if line.startswith('## '):
            text = line[3:].strip()
            add_heading(doc, text, 1)
            i += 1
            continue

        if line.startswith('### '):
            text = line[4:].strip()
            add_heading(doc, text, 2)
            i += 1
            continue

        if line.startswith('#### '):
            text = line[5:].strip()
            add_heading(doc, text, 3)
            i += 1
            continue

        # Horizontal rules
        if line.strip() == '---':
            doc.add_paragraph('─' * 60)
            i += 1
            continue

        # Bold lines
        if line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip()[2:-2]
            add_paragraph(doc, text, bold=True)
            i += 1
            continue

        # List items
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            # Handle bold within list
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for idx, part in enumerate(parts):
                if idx % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                    run.font.size = Pt(11)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(11)
            i += 1
            continue

        # Numbered list
        match = re.match(r'^\s*(\d+)\.\s+(.*)', line)
        if match:
            text = match.group(2).strip()
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for idx, part in enumerate(parts):
                if idx % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                    run.font.size = Pt(11)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(11)
            i += 1
            continue

        # Empty lines
        if line.strip() == '':
            i += 1
            continue

        # Regular paragraphs
        text = line.strip()
        if text:
            p = doc.add_paragraph()
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for idx, part in enumerate(parts):
                if idx % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                    run.font.size = Pt(11)
                else:
                    # Handle italic
                    italic_parts = re.split(r'\*(.*?)\*', part)
                    for iidx, ipart in enumerate(italic_parts):
                        run = p.add_run(ipart)
                        run.italic = (iidx % 2 == 1)
                        run.font.size = Pt(11)

        i += 1

    doc.save(docx_path)
    print(f"Word document saved to: {docx_path}")

if __name__ == '__main__':
    process_md_file('/workspace/docs/THESIS_REPORT.md', '/workspace/docs/THESIS_REPORT.docx')